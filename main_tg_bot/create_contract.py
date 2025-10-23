import os
import re
from datetime import datetime

import docx
import pythoncom
from docx2pdf import convert
from num2words import num2words
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    CommandHandler,
    ConversationHandler,
    MessageHandler,
    filters,
    ContextTypes,
    CallbackQueryHandler
)

from common.logging_config import setup_logger

logger = setup_logger("create_contract")

# Шаблоны документов
TEMPLATES = {
  "HALO Title": {
    "contract": "site_plugin/HALO_Title_contract.docx",
    "confirmation": "site_plugin/HALO_Title_confirmation.docx"
  },
  "Citygate Р311": {
    "contract": "site_plugin/Citygate_P311_contract.docx",
    "confirmation": "site_plugin/Citygate_P311_confirmation.docx"
  }
}


def amount_to_words(amount):
  """Конвертирует сумму в пропись"""
  try:
    num = re.search(r'\d+', str(amount).replace(' ', '').group())
    words = num2words(int(num), lang='ru')
    return words.capitalize()
  except:
    return str(amount)


def generate_filename(context, doc_type):
  """Генерирует имя файла без расширения"""
  object_short = re.sub(r'[^a-zA-Zа-яА-Я0-9]', '',
                        context.user_data['contract_object'])
  last_name = context.user_data['fullname'].split()[0]
  doc_date = datetime.now().strftime("%d.%m.%Y")

  if doc_type == "contract":
    return f"Договор_найма_{object_short}_краткосрок_{doc_date}_{last_name}"
  else:
    return f"Подтверждение_{object_short}_{doc_date}_{last_name}"


def process_template(template_path, replacements):
  """Обрабатывает шаблон с заменой всех полей"""
  doc = docx.Document(template_path)

  # Все возможные варианты названий полей
  field_variants = {
    'TOTAL_AMOUNT': ['TOTAL_AMOUNT', 'AMOUNT_TOTAL', 'TOTAL'],
    'TOTAL_AMOUNT_WORDS': ['TOTAL_AMOUNT_WORDS', 'AMOUNT_WORDS', 'TOTAL_WORDS'],
    'PREPAYMENT': ['PREPAYMENT', 'DEPOSIT', 'ADVANCE_PAYMENT'],
    'PREPAYMENT_BATH': ['PREPAYMENT_BATH', 'DEPOSIT_BATH', 'ADVANCE_BATH'],
    'PREPAYMENT_RUB': ['PREPAYMENT_RUB', 'DEPOSIT_RUB', 'ADVANCE_RUB']
  }

  # Обработка параграфов
  for paragraph in doc.paragraphs:
    for key, value in replacements.items():
      clean_key = key.strip('{}')
      if clean_key in field_variants:
        for variant in field_variants[clean_key]:
          if f"{{{variant}}}" in paragraph.text:
            paragraph.text = paragraph.text.replace(f"{{{variant}}}", value)
      elif key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)

  # Обработка таблиц
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for key, value in replacements.items():
          clean_key = key.strip('{}')
          if clean_key in field_variants:
            for variant in field_variants[clean_key]:
              if f"{{{variant}}}" in cell.text:
                cell.text = cell.text.replace(f"{{{variant}}}", value)
          elif key in cell.text:
            cell.text = cell.text.replace(key, value)

  return doc


async def start_contract(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Начало создания договора"""
  keyboard = [
    [InlineKeyboardButton("HALO Title", callback_data="HALO Title")],
    [InlineKeyboardButton("Citygate Р311", callback_data="Citygate Р311")]
  ]
  await update.message.reply_text(
      "Выберите объект для договора:",
      reply_markup=InlineKeyboardMarkup(keyboard))
  return SELECT_OBJECT


async def select_object(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка выбора объекта"""
  query = update.callback_query
  await query.answer()

  selected = query.data
  if selected not in TEMPLATES:
    await query.edit_message_text("Неверный выбор объекта")
    return SELECT_OBJECT

  context.user_data['contract_object'] = selected
  await query.edit_message_text(
      f"Выбран объект: {selected}\n"
      "Введите ФИО арендатора полностью:")
  return ENTER_FULLNAME


async def enter_fullname(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода ФИО"""
  context.user_data['fullname'] = update.message.text
  await update.message.reply_text("Введите серию паспорта (4 цифры):")
  return ENTER_PASSPORT_SERIES


async def enter_passport_series(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода серии паспорта"""
  if not re.match(r'^\d{4}$', update.message.text.replace(" ", "")):
    await update.message.reply_text("Неверный формат. Введите 4 цифры:")
    return ENTER_PASSPORT_SERIES

  context.user_data['passport_series'] = update.message.text.replace(" ", "")
  await update.message.reply_text("Введите номер паспорта (6 цифр):")
  return ENTER_PASSPORT_NUMBER


async def enter_passport_number(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода номера паспорта"""
  if not re.match(r'^\d{6}$', update.message.text):
    await update.message.reply_text("Неверный формат. Введите 6 цифр:")
    return ENTER_PASSPORT_NUMBER

  context.user_data['passport_number'] = update.message.text
  await update.message.reply_text("Введите, кем выдан паспорт:")
  return ENTER_PASSPORT_ISSUED


async def enter_passport_issued(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода органа выдачи"""
  context.user_data['passport_issued'] = update.message.text
  await update.message.reply_text("Введите дату выдачи (ДД.ММ.ГГГГ):")
  return ENTER_PASSPORT_DATE


async def enter_passport_date(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода даты выдачи"""
  try:
    datetime.strptime(update.message.text, "%d.%m.%Y")
    context.user_data['passport_date'] = update.message.text
    await update.message.reply_text("Введите телефон:")
    return ENTER_PHONE
  except ValueError:
    await update.message.reply_text("Неверный формат. Введите ДД.ММ.ГГГГ:")
    return ENTER_PASSPORT_DATE


async def enter_phone(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода телефона"""
  if not re.match(r'^\+?[\d\s\-\(\)]{7,}$', update.message.text):
    await update.message.reply_text("Неверный формат телефона. Попробуйте еще:")
    return ENTER_PHONE

  context.user_data['phone'] = update.message.text
  await update.message.reply_text("Введите общую сумму в батах:")
  return ENTER_TOTAL_AMOUNT


async def enter_total_amount(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода суммы"""
  if not re.match(r'^\d+$', update.message.text):
    await update.message.reply_text("Введите только цифры:")
    return ENTER_TOTAL_AMOUNT

  amount = int(update.message.text)
  context.user_data['total_amount_num'] = str(amount)
  context.user_data['total_amount'] = f"{amount} батт"
  context.user_data['total_amount_words'] = amount_to_words(amount)
  await update.message.reply_text("Введите предоплату в батах:")
  return ENTER_PREPAYMENT_BATH


async def enter_prepayment_bath(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода предоплаты в батах"""
  if not re.match(r'^\d+$', update.message.text):
    await update.message.reply_text("Введите только цифры:")
    return ENTER_PREPAYMENT_BATH

  context.user_data['prepayment_bath'] = update.message.text
  context.user_data['prepayment_bath_words'] = amount_to_words(
    update.message.text)
  await update.message.reply_text("Введите предоплату в рублях:")
  return ENTER_PREPAYMENT_RUB


async def enter_prepayment_rub(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода предоплаты в рублях"""
  if not re.match(r'^\d+$', update.message.text):
    await update.message.reply_text("Введите только цифры:")
    return ENTER_PREPAYMENT_RUB

  rub = int(update.message.text)
  bath = int(context.user_data['prepayment_bath'])
  context.user_data['prepayment_rub'] = str(rub)
  context.user_data['prepayment_rub_words'] = amount_to_words(rub)
  context.user_data['prepayment'] = (
    f"{bath} батт ({amount_to_words(bath)} батт) / "
    f"{rub} руб. ({amount_to_words(rub)} рублей)"
  )
  await update.message.reply_text("Введите дату заселения (ДД.ММ.ГГГГ):")
  return ENTER_CHECK_IN


async def enter_check_in(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода даты заселения"""
  try:
    datetime.strptime(update.message.text, "%d.%m.%Y")
    context.user_data['check_in'] = update.message.text
    await update.message.reply_text("Введите дату выезда (ДД.ММ.ГГГГ):")
    return ENTER_CHECK_OUT
  except ValueError:
    await update.message.reply_text("Неверный формат. Введите ДД.ММ.ГГГГ:")
    return ENTER_CHECK_IN


async def enter_check_out(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Обработка ввода даты выезда"""
  try:
    check_out = datetime.strptime(update.message.text, "%d.%m.%Y")
    check_in = datetime.strptime(context.user_data['check_in'], "%d.%m.%Y")

    if check_out <= check_in:
      await update.message.reply_text(
          "Дата выезда должна быть позже даты заселения:")
      return ENTER_CHECK_OUT

    context.user_data['check_out'] = update.message.text
    context.user_data['days'] = str((check_out - check_in).days)

    # Формируем сводку
    summary = (
      f"<b>Проверьте данные:</b>\n\n"
      f"<b>Объект:</b> {context.user_data['contract_object']}\n"
      f"<b>ФИО:</b> {context.user_data['fullname']}\n"
      f"<b>Паспорт:</b> {context.user_data['passport_series']} {context.user_data['passport_number']}\n"
      f"<b>Выдан:</b> {context.user_data['passport_issued']}, {context.user_data['passport_date']}\n"
      f"<b>Телефон:</b> {context.user_data['phone']}\n"
      f"<b>Даты аренды:</b> {context.user_data['check_in']} - {context.user_data['check_out']} ({context.user_data['days']} дней)\n"
      f"<b>Сумма:</b> {context.user_data['total_amount']} ({context.user_data['total_amount_words']} батт)\n"
      f"<b>Предоплата:</b> {context.user_data['prepayment']}\n\n"
      f"<b>Все верно?</b>"
    )

    keyboard = [
      [InlineKeyboardButton("✅ Да", callback_data="confirm_yes")],
      [InlineKeyboardButton("❌ Нет", callback_data="confirm_no")]
    ]
    await update.message.reply_text(
        summary,
        parse_mode='HTML',
        reply_markup=InlineKeyboardMarkup(keyboard))
    return CONFIRM_DATA

  except ValueError:
    await update.message.reply_text("Неверный формат. Введите ДД.ММ.ГГГГ:")
    return ENTER_CHECK_OUT


async def confirm_data(update: Update,
    context: ContextTypes.DEFAULT_TYPE) -> int:
  """Генерация и отправка PDF-документов"""
  query = update.callback_query
  await query.answer()

  if query.data != "confirm_yes":
    await query.edit_message_text("Создание отменено")
    return ConversationHandler.END

  try:
    pythoncom.CoInitialize()

    # Подготовка данных для шаблонов
    replacements = {
      "{FULLNAME}": context.user_data['fullname'],
      "{PASSPORT_SERIES}": context.user_data['passport_series'],
      "{PASSPORT_NUMBER}": context.user_data['passport_number'],
      "{PASSPORT_ISSUED}": context.user_data['passport_issued'],
      "{PASSPORT_DATE}": context.user_data['passport_date'],
      "{PHONE}": context.user_data['phone'],
      "{OBJECT}": context.user_data['contract_object'],

      # Денежные поля
      "{TOTAL_AMOUNT}": context.user_data['total_amount'],
      "{TOTAL_AMOUNT_NUM}": context.user_data['total_amount_num'],
      "{TOTAL_AMOUNT_WORDS}": context.user_data['total_amount_words'],

      "{PREPAYMENT}": context.user_data['prepayment'],
      "{PREPAYMENT_BATH}": context.user_data['prepayment_bath'],
      "{PREPAYMENT_BATH_WORDS}": context.user_data['prepayment_bath_words'],
      "{PREPAYMENT_RUB}": context.user_data['prepayment_rub'],
      "{PREPAYMENT_RUB_WORDS}": context.user_data['prepayment_rub_words'],

      "{CHECK_IN}": context.user_data['check_in'],
      "{CHECK_OUT}": context.user_data['check_out'],
      "{DAYS}": context.user_data['days'],
      "{TODAY}": datetime.now().strftime("%d.%m.%Y"),

      # Альтернативные названия полей
      "{AMOUNT_TOTAL}": context.user_data['total_amount'],
      "{AMOUNT_TOTAL_WORDS}": context.user_data['total_amount_words'],
      "{DEPOSIT}": context.user_data['prepayment'],
      "{DEPOSIT_BATH}": context.user_data['prepayment_bath'],
      "{DEPOSIT_RUB}": context.user_data['prepayment_rub']
    }

    # Создаем папку для документов
    os.makedirs("generated_contracts", exist_ok=True)

    # Генерация договора
    contract_name = generate_filename(context, "contract")
    contract_docx_path = os.path.join("generated_contracts",
                                      f"{contract_name}.docx")
    contract_pdf_path = os.path.join("generated_contracts",
                                     f"{contract_name}.pdf")

    contract = process_template(
        TEMPLATES[context.user_data['contract_object']]["contract"],
        replacements
    )
    contract.save(contract_docx_path)

    # Конвертация в PDF
    convert(contract_docx_path, contract_pdf_path)

    # Генерация подтверждения
    confirm_name = generate_filename(context, "confirmation")
    confirm_docx_path = os.path.join("generated_contracts",
                                     f"{confirm_name}.docx")
    confirm_pdf_path = os.path.join("generated_contracts",
                                    f"{confirm_name}.pdf")

    confirmation = process_template(
        TEMPLATES[context.user_data['contract_object']]["confirmation"],
        replacements
    )
    confirmation.save(confirm_docx_path)

    # Конвертация в PDF
    convert(confirm_docx_path, confirm_pdf_path)

    # Отправка PDF-документов
    with open(contract_pdf_path, 'rb') as f1, open(confirm_pdf_path,
                                                   'rb') as f2:
      await query.message.reply_document(f1, caption="Договор найма (PDF)")
      await query.message.reply_document(f2,
                                         caption="Подтверждение бронирования (PDF)")

    # Удаляем временные DOCX-файлы
    os.remove(contract_docx_path)
    os.remove(confirm_docx_path)

    logger.info(
      f"Generated PDF documents: {contract_pdf_path}, {confirm_pdf_path}")

  except Exception as e:
    logger.error(f"Document generation failed: {str(e)}", exc_info=True)
    await query.message.reply_text(
      f"⚠️ Ошибка при создании документов: {str(e)}")

  finally:
    pythoncom.CoUninitialize()

  return ConversationHandler.END


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
  """Отмена создания"""
  await update.message.reply_text("Создание договора отменено")
  return ConversationHandler.END


def get_contract_conversation_handler() -> ConversationHandler:
  """Фабрика обработчика диалога"""
  return ConversationHandler(
      entry_points=[CommandHandler("create_contract", start_contract)],
      states={
        SELECT_OBJECT: [CallbackQueryHandler(select_object)],
        ENTER_FULLNAME: [
          MessageHandler(filters.TEXT & ~filters.COMMAND, enter_fullname)],
        ENTER_PASSPORT_SERIES: [MessageHandler(filters.TEXT & ~filters.COMMAND,
                                               enter_passport_series)],
        ENTER_PASSPORT_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND,
                                               enter_passport_number)],
        ENTER_PASSPORT_ISSUED: [MessageHandler(filters.TEXT & ~filters.COMMAND,
                                               enter_passport_issued)],
        ENTER_PASSPORT_DATE: [
          MessageHandler(filters.TEXT & ~filters.COMMAND, enter_passport_date)],
        ENTER_PHONE: [
          MessageHandler(filters.TEXT & ~filters.COMMAND, enter_phone)],
        ENTER_TOTAL_AMOUNT: [
          MessageHandler(filters.TEXT & ~filters.COMMAND, enter_total_amount)],
        ENTER_PREPAYMENT_BATH: [MessageHandler(filters.TEXT & ~filters.COMMAND,
                                               enter_prepayment_bath)],
        ENTER_PREPAYMENT_RUB: [MessageHandler(filters.TEXT & ~filters.COMMAND,
                                              enter_prepayment_rub)],
        ENTER_CHECK_IN: [
          MessageHandler(filters.TEXT & ~filters.COMMAND, enter_check_in)],
        ENTER_CHECK_OUT: [
          MessageHandler(filters.TEXT & ~filters.COMMAND, enter_check_out)],
        CONFIRM_DATA: [CallbackQueryHandler(confirm_data)],
      },
      fallbacks=[CommandHandler("cancel", cancel)],
  )