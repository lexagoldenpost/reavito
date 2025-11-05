# main_tg_bot/handlers/contract_handler.py

import pandas as pd
from pathlib import Path
from typing import Any, Dict, Optional
from docx2pdf import convert
import docx
from datetime import datetime
import uuid
import aiohttp
import os
import tempfile

from num2words import num2words

from common.logging_config import setup_logger
from main_tg_bot.sender.tg_notifier import send_message

logger = setup_logger("contract_handler")

# Корень проекта — родитель main_tg_bot/
PROJECT_ROOT = Path(__file__).parent.parent.resolve()
# Константы для путей к шаблонам
TEMPLATE_DIR = PROJECT_ROOT / "word_templates"
CONTRACT_TEMPLATE_PREFIX = "Договор"
CONFIRMATION_TEMPLATE_PREFIX = "Подтверждение"


def format_number_with_spaces(value: str) -> str:
    """
    Принимает строку, представляющую целое число, и возвращает её с пробелами как разделителями тысяч.
    Пример: "1000000" → "1 000 000"
    Если не удаётся преобразовать — возвращает исходную строку.
    """
    try:
        num = int(value)
        return f"{num:,}".replace(",", " ")
    except (ValueError, TypeError):
        return value


async def handle_contract(data: Dict[str, Any], filename: str):
    """
    Обработчик генерации договоров и подтверждений
    """
    logger.info("📄 [contract_handler] Начало обработки договора")
    logger.info(f"📄 [contract_handler] Имя файла: {filename}")
    logger.info(f"📄 [contract_handler] Данные договора:")
    for key, value in data.items():
        logger.info(f"    {key}: {value}")

    init_chat_id: Optional[str] = data.get('init_chat_id')
    guest_name: str = data.get('fullname', '').strip()

    # --- Сразу отправляем "обрабатывается" ---
    if init_chat_id:
        try:
            async with aiohttp.ClientSession() as session:
                await send_message(session, init_chat_id, f"📄 Договора и подтверждение {guest_name} формируется, ожидайте...")
                logger.info(f"📢 Уведомление 'обрабатывается' отправлено в чат {init_chat_id}")
        except Exception as e:
            logger.warning(f"Не удалось отправить начальное уведомление в Telegram: {e}")

    try:
        # --- Валидация обязательных полей ---
        required_fields = [
            'contract_object', 'contract_type', 'fullname',
            'passport_series', 'passport_number', 'passport_issued',
            'passport_date', 'phone', 'check_in', 'check_out',
            'total_amount', 'prepayment_bath'
        ]

        missing_fields = [field for field in required_fields if not data.get(field)]
        if missing_fields:
            raise ValueError(f"❌ Отсутствуют обязательные поля: {', '.join(missing_fields)}")

        # --- Извлечение номера документа из имени файла ---
        if not filename.endswith('.json'):
            raise ValueError(f"❌ Имя файла должно заканчиваться на .json, получено: {filename}")

        # Убираем .json
        base_name = filename[:-5]

        # Проверяем, что имя начинается с "Договор_"
        if not base_name.startswith("Договор_"):
            raise ValueError(f"❌ Имя файла должно начинаться с 'Договор_', получено: {base_name}")

        # Формируем номер подтверждения: заменяем "Договор" на "Подтверждение"
        confirmation_base_name = "Подтверждение" + base_name[len("Договор"):]

        contract_number = base_name
        confirmation_number = confirmation_base_name

        logger.info(f"📄 Номер договора из файла: {contract_number}")
        logger.info(f"📄 Номер подтверждения: {confirmation_number}")

        # --- Подготовка данных для шаблонов ---
        contract_data = prepare_template_data(data, contract_number)

        # --- Подготовка данных для подтверждения ---
        confirmation_data = prepare_template_data(data, confirmation_number)

        # --- Определение путей к шаблонам ---
        contract_template_path = TEMPLATE_DIR / f"{CONTRACT_TEMPLATE_PREFIX}_{data['contract_object']}_{data['contract_type']}.docx"
        confirmation_template_path = TEMPLATE_DIR / f"{CONFIRMATION_TEMPLATE_PREFIX}_{data['contract_object']}_{data['contract_type']}.docx"

        logger.info(f"📄 Путь к шаблону договора: {contract_template_path}")
        logger.info(f"📄 Путь к шаблону подтверждения: {confirmation_template_path}")

        # --- Проверка существования шаблонов ---
        if not contract_template_path.exists():
            raise FileNotFoundError(f"❌ Шаблон договора не найден: {contract_template_path}")
        if not confirmation_template_path.exists():
            raise FileNotFoundError(f"❌ Шаблон подтверждения не найден: {confirmation_template_path}")

        # --- Генерация документов ---
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)

            # Генерация договора
            contract_docx_path = temp_path / f"contract_{contract_number}.docx"
            contract_pdf_path = temp_path / f"contract_{contract_number}.pdf"

            await fill_template(contract_template_path, contract_docx_path, contract_data)
            await convert_to_pdf(contract_docx_path, contract_pdf_path)

            # Генерация подтверждения
            confirmation_docx_path = temp_path / f"confirmation_{contract_number}.docx"
            confirmation_pdf_path = temp_path / f"confirmation_{contract_number}.pdf"

            await fill_template(confirmation_template_path, confirmation_docx_path, confirmation_data)
            await convert_to_pdf(confirmation_docx_path, confirmation_pdf_path)

            # --- Отправка файлов пользователю ---
            if init_chat_id:
                async with aiohttp.ClientSession() as session:
                    # Отправка договора
                    with open(contract_pdf_path, 'rb') as contract_file:
                        await send_message(
                            session,
                            init_chat_id,
                            f"📄 Договор аренды для {data['fullname']}",
                            document=contract_file,
                            filename=f"{contract_number}.pdf"
                        )

                    # Отправка подтверждения
                    with open(confirmation_pdf_path, 'rb') as confirmation_file:
                        await send_message(
                            session,
                            init_chat_id,
                            f"✅ Подтверждение бронирования для {data['fullname']}",
                            document=confirmation_file,
                            filename=f"{confirmation_number}.pdf"
                        )
                    success_msg = f"✅ Документы для {data['fullname']} успешно сгенерированы и отправлены!"
                    await send_message(session, init_chat_id, success_msg)
                    logger.info(f"✅ Уведомление об успешной генерации отправлено в чат {init_chat_id}")

        logger.info("📄 [contract_handler] Генерация договоров завершена успешно")

    except Exception as e:
        error_msg = str(e)
        logger.error(f"❌ Ошибка при генерации договора: {error_msg}")
        if init_chat_id:
            async with aiohttp.ClientSession() as session:
                await send_message(
                    session,
                    init_chat_id,
                    f"❌ Произошла ошибка при генерации договора: {error_msg}"
                )


def prepare_template_data(data: Dict[str, Any], contract_number: str) -> Dict[str, str]:
    """
    Подготовка данных для заполнения шаблонов, включая суммы прописью и форматированные числа
    """
    def bath_to_words(amount: str) -> str:
        try:
            value = int(amount)
            words = num2words(value, lang='en').capitalize()
            return f"{words} Baht"
        except (ValueError, TypeError):
            return ""

    def rub_to_words(amount: str) -> str:
        try:
            value = int(amount)
            words = num2words(value, lang='ru', to='currency', currency='RUB')
            return words.capitalize()
        except (ValueError, TypeError):
            return ""

    # Исходные значения
    total_amount_raw = data.get('total_amount', '0')
    prepayment_bath_raw = data.get('prepayment_bath', '0')
    prepayment_rub_raw = data.get('prepayment_rub', '0')

    # Форматированные значения с пробелами
    total_amount = format_number_with_spaces(total_amount_raw)
    prepayment_bath = format_number_with_spaces(prepayment_bath_raw)
    prepayment_rub = format_number_with_spaces(prepayment_rub_raw)

    # Расчёт остатка в батах
    try:
        total_int = int(total_amount_raw)
        prepayment_int = int(prepayment_bath_raw)
        final_payment_bath_raw = total_int - prepayment_int
        final_payment_bath = format_number_with_spaces(str(final_payment_bath_raw))
        final_payment_bath_words_th = bath_to_words(str(final_payment_bath_raw))
    except (ValueError, TypeError):
        final_payment_bath = total_amount
        final_payment_bath_words_th = ""

    # Суммы прописью (на основе исходных чисел)
    total_amount_words_th = bath_to_words(total_amount_raw)
    prepayment_bath_words_th = bath_to_words(prepayment_bath_raw)
    prepayment_rub_words_ru = rub_to_words(prepayment_rub_raw)

    template_data = {
        # Основная информация
        'contract_number': contract_number,
        'current_date': datetime.now().strftime("%d.%m.%Y"),

        # Данные клиента
        'fullname': data.get('fullname', ''),
        'passport_series': data.get('passport_series', ''),
        'passport_number': data.get('passport_number', ''),
        'passport_full': f"{data.get('passport_series', '')} {data.get('passport_number', '')}",
        'passport_issued': data.get('passport_issued', ''),
        'passport_date': data.get('passport_date', ''),
        'phone': data.get('phone', ''),

        # Даты бронирования
        'check_in': data.get('check_in', ''),
        'check_out': data.get('check_out', ''),

        # Финансовые данные (цифрами, с пробелами)
        'total_amount': total_amount,
        'prepayment_bath': prepayment_bath,
        'prepayment_rub': prepayment_rub,
        'final_payment_bath': final_payment_bath,  # ← новое поле

        # Финансовые данные (прописью)
        'total_amount_words_th': total_amount_words_th,
        'prepayment_bath_words_th': prepayment_bath_words_th,
        'prepayment_rub_words_ru': prepayment_rub_words_ru,
        'final_payment_bath_words_th': final_payment_bath_words_th,  # ← новое поле

        # Объект недвижимости
        'contract_object': data.get('contract_object', ''),
        'contract_type': data.get('contract_type', ''),

        # Текущий год
        'current_year': datetime.now().strftime("%Y")
    }

    logger.info(f"📄 Подготовлены данные для шаблона: {template_data}")
    return template_data

async def fill_template(template_path: Path, output_path: Path, data: Dict[str, str]):
    """
    Заполнение шаблона DOCX данными
    """
    try:
        doc = docx.Document(template_path)

        # Замена текста в параграфах
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f'{{{key}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{key}}}', value)

        # Замена текста в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f'{{{key}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{key}}}', value)

        doc.save(output_path)
        logger.info(f"✅ Шаблон заполнен: {output_path}")

    except Exception as e:
        logger.error(f"❌ Ошибка при заполнении шаблона {template_path}: {e}")
        raise


async def convert_to_pdf(docx_path: Path, pdf_path: Path):
    """
    Конвертация DOCX в PDF
    """
    try:
        convert(docx_path, pdf_path)
        logger.info(f"✅ DOCX конвертирован в PDF: {pdf_path}")
    except Exception as e:
        logger.error(f"❌ Ошибка при конвертации {docx_path} в PDF: {e}")
        raise